import streamlit as st
from docx import Document
from deep_translator import GoogleTranslator
import spacy
import openai
nlp = spacy.load("en_core_web_sm")
openai.api_key = "Enter your api key"
import io


st.set_page_config(page_icon="📄",page_title='Translator')


def extract_names__(text):
    doc = nlp(text)
    extracted_names = []
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            extracted_names.append(ent.text)
    return extracted_names

def get_names(input_path,max_chars=5000):
    doc = Document(input_path)
    names = Document()
    
    current_chunk = ""
    
    for paragraph in doc.paragraphs:
        if len(current_chunk) + len(paragraph.text) > max_chars:
            extracted_names = extract_names__(current_chunk)
            names.add_paragraph(extracted_names)
            current_chunk = paragraph.text
        else:
            current_chunk += paragraph.text
    
    # Translate the last chunk
    if current_chunk:
        extract_names = extract_names__(current_chunk)
        names.add_paragraph(extract_names)
    
    return names

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text.strip()


def translate_text(text, target_language):
    translator = GoogleTranslator(source='auto', target=target_language)
    translated_text = translator.translate(text)
    return translated_text

def translate_large_text(input_text, target_language, max_chars=5000):
    translated_text = ""
    current_chunk = ""
    
    paragraphs = input_text.split("\n")  # Split input into paragraphs
    
    for paragraph in paragraphs:
        if len(current_chunk) + len(paragraph) > max_chars:
            translated_chunk = translate_text(current_chunk, target_language)
            translated_text += translated_chunk + "\n"
            current_chunk = paragraph
        else:
            current_chunk += paragraph + "\n"
    
    # Translate the last chunk
    if current_chunk:
        translated_chunk = translate_text(current_chunk, target_language)
        translated_text += translated_chunk
    
    return translated_text

def extract_names_openai(text):
    prompt = f"Extract names from the following text:\n'{text}'\n\nExtracted names:"
    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=100
    )
    extracted_text = response.choices[0].text.strip()
    return extracted_text

def localize_names(text, target_country):
    prompt = f"Extract names from the following text:\n'{text}'\n\n"
    prompt = f"Generate localized names of persons from the  {target_country}"
    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=150
    )
    localized_names = response.choices[0].text.strip()
    return localized_names

def replace_names(text, names_to_replace, replacement_names):
    for name_to_replace, replacement_name in zip(names_to_replace, replacement_names):
        text = text.replace(name_to_replace, replacement_name)
    return text

def download_button(file_data, file_name):
    # Create a download link using the st.download_button() function
    st.download_button(
        label="Download",
        data=file_data,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# Streamlit app
st.title("DOCX Translation Bot 🤖")


languages = {
    "English": "en",
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Italian": "it",
    "Dutch": "nl",
    "Portuguese": "pt",
    "Russian": "ru",
    "Chinese (Simplified)": "zh",
    "Chinese (Traditional)": "zh-TW",
    "Japanese": "ja",
    "Korean": "ko",
    "Arabic": "ar",
    "Hindi": "hi",
    "Turkish": "tr",
    "Greek": "el",
    "Swedish": "sv",
    "Norwegian": "no",
    "Danish": "da",
    "Finnish": "fi",
    "Polish": "pl",
    "Romanian": "ro",
    "Czech": "cs",
    "Thai": "th",
    "Indonesian": "id",
    "Malay": "ms",
    "Vietnamese": "vi",
    "Hebrew": "he",
    "Bengali": "bn",
    "Slovak": "sk",
    "Ukrainian": "uk",
    "Slovenian": "sl",
    "Croatian": "hr",
    "Serbian": "sr",
    "Macedonian": "mk",
    "Bulgarian": "bg",
    "Latvian": "lv",
    "Lithuanian": "lt",
    "Estonian": "et",
    "Icelandic": "is",
    "Georgian": "ka",
    "Irish": "ga",
    "Welsh": "cy",
    "Frisian": "fy",
    "Albanian": "sq",
    "Maltese": "mt",
    "Basque": "eu",
    "Catalan": "ca",
    "Galician": "gl",
    "Luxembourgish": "lb",
    "Belarusian": "be",
    "Bosnian": "bs",
    "Montenegrin": "me",
    "Armenian": "hy",
    "Azerbaijani": "az",
    "Kazakh": "kk",
    "Uzbek": "uz",
    "Tajik": "tg",
    "Kyrgyz": "ky",
    "Turkmen": "tk",
    "Uighur": "ug",
    "Tatar": "tt",
    "Bashkir": "ba",
    "Chuvash": "cv",
    "Chechen": "ce",
    "Ingush": "inh",
    "Abkhaz": "ab",
    "Ossetian": "os",
    "Dari": "prs",
    "Pashto": "ps",
    "Kurdish": "ku",
    "Baluchi": "bal",
    "Sindhi": "sd",
    "Tamil": "ta",
    "Telugu": "te",
    "Kannada": "kn",
    "Malayalam": "ml",
    "Marathi": "mr",
    "Gujarati": "gu",
    "Punjabi": "pa",
    "Oriya": "or",
    "Assamese": "as",
    "Maithili": "mai",
    "Nepali": "ne",
    "Sinhalese": "si",
    "Burmese": "my",
    "Khmer": "km",
    "Lao": "lo",
    "Thai": "th",
    "Vietnamese": "vi",
    "Indonesian": "id",
    "Tagalog": "tl",
    "Malay": "ms",
    "Javanese": "jv",
    "Sundanese": "su",
    "Madurese": "mad",
    "Minangkabau": "min",
    "Banjar": "bjn",
    "Buginese": "bug",
    "Acehnese": "ace",
    "Makassar": "mak",
    "Balinese": "ban",
    "Sasak": "sas",
    "Tetum": "tet",
    "Rotuman": "rtm",
    "Hawaiian": "haw",
    "Samoan": "sm",
    "Tongan": "to",
    "Fijian": "fj",
    "Marshallese": "mh",
    "Palauan": "pau",
    "Chamorro": "ch",
    "Pohnpeian": "pon",
    "Yapese": "yap",
    "Chuukese": "chk",
    "Kosraean": "kos",
    "Nauruan": "na",
    "Tuvaluan": "tvl",
    "Kiribati": "gil",
    "Marshall": "mah",
    "Tok Pisin": "tpi",
    "Hiri Motu": "hmo",
    "Niuatoputapu": "niu",
    "Futunan": "fud",
    "Wallisian": "wls",
    "Drehu": "dhv",
    "Iban": "iba",
    "Cebuano": "ceb",
    "Hiligaynon": "hil",
    "Waray": "war",
    "Bikol": "bik",
}
countries = [
    "India",
    "United States",
    "China",
    "Russia",
    "Brazil",
    "Japan",
    "Germany",
    "United Kingdom",
    "France",
    "Italy"
]


uploaded_file = st.file_uploader("Please upload your DOCX file Here", type="docx")
if uploaded_file is not None:
    target_language = st.selectbox("Please select your language for translation",options=list(languages.keys()))
    countries = st.selectbox("Please select the country",options=countries)
    file_name = st.text_input("Please enter your file name")
    for k,v in languages.items():
        if k == target_language:
            target_language = v
    if st.button("Translate"):
        try:
            text_doc = extract_text_from_docx(uploaded_file)
            extract_names_ = get_names(uploaded_file)
            st.write()
            # Get the translated text directly from the translated_doc
            #translated_text = '\n'.join([para.text for para in translated_doc.paragraphs])
            extracted_names = '\n'.join([para.text for para in extract_names_.paragraphs])
            extracted_name_openai =  extract_names_openai(extracted_names)
            changed_names = localize_names(extracted_name_openai,countries)
            extracted_name_openai_list_2d = [name.split() for name in extracted_name_openai.split(", ")]
            changed_names_list_2d = [name.split() for name in changed_names.split(", ")]
            extracted_name_openai_list_1d = []
            for sublist in extracted_name_openai_list_2d:
                for item in sublist:
                    extracted_name_openai_list_1d.append(item)
            changed_names_list_1d = []
            for sublist in changed_names_list_2d:
                for item in sublist:
                    changed_names_list_1d.append(item)
            st.write(changed_names_list_1d)
            replace_text = replace_names(text_doc,extracted_name_openai_list_1d,changed_names_list_1d)
            
            translated_text = translate_large_text(replace_text,target_language)
            
            translate_doc =  Document()
            translate_doc.add_paragraph(translated_text)

            try:
                text_doc = extract_text_from_docx(uploaded_file)
                extract_names_ = get_names(uploaded_file)
        # ... (rest of your processing code)

                translate_doc = Document()
                translate_doc.add_paragraph(translated_text)

        # Save the translated document to a BytesIO buffer
                translated_doc_buffer = io.BytesIO()
                translate_doc.save(translated_doc_buffer)
                translated_doc_buffer.seek(0)

        # Call the download_button function with the BytesIO buffer
                download_button(translated_doc_buffer, f"translated_{file_name}.docx")

            except Exception as e:

                st.error(f"An error occurred: {e}")


            
            
            
        except Exception as e:
            st.error(f"An error occurred: {e}")
